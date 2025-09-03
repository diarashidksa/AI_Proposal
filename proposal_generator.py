import os
import json
import uuid
import pickle
from langdetect import detect
from pypdf import PdfReader
from docx import Document
from sentence_transformers import SentenceTransformer
import faiss
import openai

# ----------------- Config -----------------
openai.api_key = os.getenv("OPENAI_API_KEY")
TRAINING_FOLDER = "Training_Data"
OUTPUT_FOLDER = "Proposals"
FAISS_INDEX_FILE = "training_index.faiss"
CHUNKS_PKL_FILE = "chunks.pkl"
CHUNK_IDS_PKL_FILE = "chunk_ids.pkl"
CHUNK_SOURCES_PKL_FILE = "chunk_sources.pkl"


# ----------------- Helpers -----------------
def read_file_content(filepath):
    text = ""
    if filepath.endswith(".txt"):
        with open(filepath, "r", encoding="utf-8") as f:
            text = f.read()
    elif filepath.endswith(".pdf"):
        reader = PdfReader(filepath)
        text = " ".join([page.extract_text() for page in reader.pages if page.extract_text()])
    elif filepath.endswith(".docx"):
        doc = Document(filepath)
        text = " ".join([p.text for p in doc.paragraphs])
    return text


def detect_language(text):
    try:
        lang_code = detect(text)
        return "ar" if lang_code == "ar" else "en"
    except:
        return "en"


# ----------------- Training / FAISS -----------------
def build_faiss_index(training_folder):
    model = SentenceTransformer('all-MiniLM-L6-v2')
    documents = []
    doc_ids = []

    if not os.path.exists(training_folder):
        os.makedirs(training_folder)

    for fname in os.listdir(training_folder):
        path = os.path.join(training_folder, fname)
        text = read_file_content(path)
        if text.strip():
            documents.append(text)
            doc_ids.append(fname)

    chunks = []
    chunk_ids = []
    chunk_sources = []

    for i, doc in enumerate(documents):
        words = doc.split()
        for j in range(0, len(words), 500):
            chunk = " ".join(words[j:j + 500])
            chunks.append(chunk)
            chunk_ids.append(f"{doc_ids[i]}_chunk{j // 500}")
            chunk_sources.append(doc_ids[i])

    if chunks:
        embeddings = model.encode(chunks)
        dim = embeddings.shape[1]
        index = faiss.IndexFlatL2(dim)
        index.add(embeddings)
        faiss.write_index(index, FAISS_INDEX_FILE)
        with open(CHUNK_IDS_PKL_FILE, "wb") as f:
            pickle.dump(chunk_ids, f)
        with open(CHUNKS_PKL_FILE, "wb") as f:
            pickle.dump(chunks, f)
        with open(CHUNK_SOURCES_PKL_FILE, "wb") as f:
            pickle.dump(chunk_sources, f)
        print(f"FAISS index built with {len(chunks)} chunks.")
    else:
        index = None
        print("No training data found. FAISS index not built.")

    return index, chunks, chunk_ids, chunk_sources, model


def load_faiss_index_and_chunks():
    # Remove old index files to force a fresh rebuild every time
    for file in [FAISS_INDEX_FILE, CHUNKS_PKL_FILE, CHUNK_IDS_PKL_FILE, CHUNK_SOURCES_PKL_FILE]:
        if os.path.exists(file):
            os.remove(file)
            print(f"Removed old index file: {file}")

    # Build a fresh FAISS index and return the necessary data
    index, chunks, chunk_ids, chunk_sources, _ = build_faiss_index(TRAINING_FOLDER)
    return index, chunks, chunk_ids, chunk_sources


# ----------------- GPT Proposal -----------------
def generate_ai_proposal(context_text, brief, lang, faiss_sources):
    # Prompt the LLM to include a 'Sources' key
    prompt = f"""
You are a professional proposal writer. Generate a complete professional proposal in {lang} language.
Include the following sections:

1. Executive Summary
2. Scope of Work
3. Technical Proposal
4. Timeline
5. Assumptions
6. Team
7. Quality
8. Financial Proposal

Return strictly as JSON (no markdown, no ```json``` blocks) in this format:

{{
"ExecutiveSummary": "...",
"Scope": "...",
"TechnicalProposal": "...",
"Timeline": "...",
"Assumptions": "...",
"Team": "...",
"Quality": "...",
"FinancialProposal": [
    {{"Item":"...","UOM":"...","Qty":1,"UnitPrice":0,"Total":0}}
],
"Sources": {{
    "ExecutiveSummary": ["source1", "source2"],
    "Scope": ["source1", "source3"],
    ...
    "FinancialProposal": ["source5"]
}}
}}

The sources must be either one of the following:
{json.dumps(faiss_sources)}
or 'LLM (Internal Knowledge)' if the information is not from the provided documents.

Client Brief:
{brief}

Reference Documents:
{context_text}
"""
    client = openai.OpenAI()
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.4,
        max_tokens=4000
    )

    result_text = response.choices[0].message.content
    try:
        if result_text.strip().startswith("```json"):
            result_text = result_text.strip()[7:-3]
        proposal_json = json.loads(result_text)
    except json.JSONDecodeError:
        proposal_json = {
            "ExecutiveSummary": result_text,
            "Scope": "",
            "TechnicalProposal": "",
            "Timeline": "",
            "Assumptions": "",
            "Team": "",
            "Quality": "",
            "FinancialProposal": [],
            "Sources": {"ExecutiveSummary": ["LLM (Internal Knowledge)"]}
        }

    # Ensure a sources key exists and fill with a default if not
    if "Sources" not in proposal_json:
        proposal_json["Sources"] = {}
        for key in ["ExecutiveSummary", "Scope", "TechnicalProposal", "Timeline", "Assumptions", "Team", "Quality"]:
            if key in proposal_json:
                proposal_json["Sources"][key] = ["LLM (Internal Knowledge)"]

    return proposal_json


# ----------------- DOCX Generation -----------------
def build_docx(proposal_json, output_folder, lang):
    doc = Document()
    doc.add_heading("Proposal", 0)

    section_titles = {
        "en": {
            "ExecutiveSummary": "Executive Summary",
            "Scope": "Scope of Work",
            "TechnicalProposal": "Technical Proposal",
            "Timeline": "Timeline",
            "Assumptions": "Assumptions",
            "Team": "Team",
            "Quality": "Quality",
            "FinancialProposal": "Financial Proposal",
            "Sources": "Resources & References"
        },
        "ar": {
            "ExecutiveSummary": "الملخص التنفيذي",
            "Scope": "نطاق العمل",
            "TechnicalProposal": "الخطة الفنية",
            "Timeline": "الجدول الزمني",
            "Assumptions": "الافتراضات",
            "Team": "الفريق",
            "Quality": "الجودة",
            "FinancialProposal": "الاقتراح المالي",
            "Sources": "المصادر والمراجع"
        }
    }

    lang_code = "ar" if lang == "ar" else "en"
    titles = section_titles[lang_code]

    proposal_content_keys = ["ExecutiveSummary", "Scope", "TechnicalProposal", "Timeline", "Assumptions", "Team",
                             "Quality"]

    for key in proposal_content_keys:
        content = proposal_json.get(key, "")
        if content.strip():
            doc.add_heading(titles[key], level=1)
            doc.add_paragraph(content.strip())

    finance_items = proposal_json.get("FinancialProposal", [])
    if finance_items:
        doc.add_heading(titles["FinancialProposal"], level=1)
        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Item" if lang_code == "en" else "البند"
        hdr_cells[1].text = "UOM" if lang_code == "en" else "الوحدة"
        hdr_cells[2].text = "Qty" if lang_code == "en" else "الكمية"
        hdr_cells[3].text = "Unit Price" if lang_code == "en" else "سعر الوحدة"
        hdr_cells[4].text = "Total" if lang_code == "en" else "الإجمالي"

        total_sum = 0
        for row in finance_items:
            cells = table.add_row().cells
            cells[0].text = str(row.get("Item", ""))
            cells[1].text = str(row.get("UOM", ""))
            cells[2].text = str(row.get("Qty", 0))
            cells[3].text = str(row.get("UnitPrice", 0))
            cells[4].text = str(row.get("Total", 0))
            total_sum += row.get("Total", 0)

        cells = table.add_row().cells
        cells[0].merge(cells[1])
        cells[0].merge(cells[2])
        cells[0].merge(cells[3])
        cells[0].text = "Grand Total" if lang_code == "en" else "الإجمالي الكلي"
        cells[4].text = str(total_sum)

    # Add the resources section
    sources_data = proposal_json.get("Sources", {})
    if sources_data:
        doc.add_heading(titles["Sources"], level=1)
        for key in proposal_content_keys + ["FinancialProposal"]:
            sources = sources_data.get(key, [])
            if sources:
                unique_sources = sorted(list(set(sources)))
                doc.add_heading(titles.get(key, key), level=2)
                for source in unique_sources:
                    doc.add_paragraph(f"- {source}")

    output_filename = f"proposal_{uuid.uuid4().hex}.docx"
    output_path = os.path.join(output_folder, output_filename)
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    doc.save(output_path)

    return output_filename, finance_items, sources_data


# ----------------- Main Proposal Generator -----------------
def generate_proposal(file_paths, brief):
    all_text = ""
    for path in file_paths:
        all_text += read_file_content(path) + "\n"

    lang = detect_language(all_text if all_text.strip() else brief)

    index, chunks, chunk_ids, chunk_sources = load_faiss_index_and_chunks()
    model = SentenceTransformer('all-MiniLM-L6-v2')

    # Create a mapping from chunk_id to source filename
    chunk_to_source_map = {chunk_id: source for chunk_id, source in zip(chunk_ids, chunk_sources)}

    user_text = all_text + "\n" + brief
    query_embedding = model.encode([user_text])

    retrieved_text = ""
    faiss_document_sources = []

    if chunks and index:
        k = min(5, len(chunks))
        distances, indices = index.search(query_embedding, k)
        for i in indices[0]:
            retrieved_text += chunks[i] + "\n"
            faiss_document_sources.append(chunk_to_source_map[chunk_ids[i]])

    context_text = retrieved_text + "\n\n" + user_text

    # Pass retrieved sources to the LLM to guide it on what to cite
    proposal_json = generate_ai_proposal(context_text, brief, lang, list(set(faiss_document_sources)))

    output_filename, finance_items, sources_data = build_docx(proposal_json, OUTPUT_FOLDER, lang)

    print(f"Proposal generated: {output_filename}")
    print("Resources used:")
    print(json.dumps(sources_data, indent=2))

    return output_filename, finance_items, sources_data